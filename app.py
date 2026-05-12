import os
import textwrap
from pathlib import Path

from flask import Flask, render_template, request, send_file, redirect, url_for, flash
from werkzeug.utils import secure_filename
from deep_translator import GoogleTranslator
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas

from pdf2image import convert_from_path
import pytesseract
from PIL import ImageOps, ImageFilter


BASE_DIR = Path(__file__).resolve().parent
UPLOAD_FOLDER = BASE_DIR / "uploads"
OUTPUT_FOLDER = BASE_DIR / "output"
TEMPLATES_FOLDER = BASE_DIR / "templates"

ALLOWED_EXTENSIONS = {".pdf", ".docx"}

UPLOAD_FOLDER.mkdir(exist_ok=True)
OUTPUT_FOLDER.mkdir(exist_ok=True)
TEMPLATES_FOLDER.mkdir(exist_ok=True)

app = Flask(__name__, template_folder=str(TEMPLATES_FOLDER))
app.secret_key = "super-secret-key-change-this"

app.config["UPLOAD_FOLDER"] = str(UPLOAD_FOLDER)
app.config["OUTPUT_FOLDER"] = str(OUTPUT_FOLDER)

TESSERACT_PATH = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
POPPLER_PATH = r"C:\Users\sanju\Downloads\Release-25.12.0-0\poppler-25.12.0\Library\bin"

if os.path.exists(TESSERACT_PATH):
    pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH

OCR_LANGUAGE_MAP = {
    "english": "eng",
    "hindi": "hin+eng",
    "gujarati": "guj+eng",
    "marathi": "mar+eng",
    "urdu": "urd+eng",
}


def allowed_file(filename: str) -> bool:
    return Path(filename).suffix.lower() in ALLOWED_EXTENSIONS


def extract_text_from_pdf(file_path: str) -> str:
    text_parts = []
    reader = PdfReader(file_path)

    for page in reader.pages:
        page_text = page.extract_text() or ""
        if page_text.strip():
            text_parts.append(page_text)

    return "\n".join(text_parts).strip()


def extract_text_from_docx(file_path: str) -> str:
    doc = Document(file_path)
    text_parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    return "\n".join(text_parts).strip()


def preprocess_image_for_ocr(img):
    img = ImageOps.grayscale(img)
    img = img.filter(ImageFilter.SHARPEN)
    img = img.point(lambda x: 0 if x < 160 else 255, "1")
    return img


def validate_ocr_language(ocr_lang: str) -> None:
    try:
        installed_langs = set(pytesseract.get_languages(config=""))
    except Exception as e:
        raise RuntimeError(f"Unable to read installed Tesseract languages. Error: {e}")

    needed_langs = set(ocr_lang.split("+"))
    missing_langs = needed_langs - installed_langs

    if missing_langs:
        missing_text = ", ".join(sorted(missing_langs))
        raise RuntimeError(
            f"Tesseract language data missing: {missing_text}. "
            f"Please add the required .traineddata files in tessdata folder."
        )


def extract_text_from_pdf_ocr(file_path: str, source_lang: str = "english") -> str:
    if not os.path.exists(TESSERACT_PATH):
        raise RuntimeError(
            f"Tesseract not found at: {TESSERACT_PATH}. "
            f"Install Tesseract OCR or fix TESSERACT_PATH."
        )

    if not os.path.exists(POPPLER_PATH):
        raise RuntimeError(
            f"Poppler not found at: {POPPLER_PATH}. "
            f"Install Poppler or fix POPPLER_PATH."
        )

    ocr_lang = OCR_LANGUAGE_MAP.get(source_lang.lower(), "eng")
    validate_ocr_language(ocr_lang)

    try:
        images = convert_from_path(
            file_path,
            dpi=300,
            poppler_path=POPPLER_PATH
        )
    except Exception as e:
        raise RuntimeError(
            f"Unable to convert PDF pages to images. Check Poppler. Error: {e}"
        )

    text_parts = []

    for img in images:
        try:
            processed_img = preprocess_image_for_ocr(img)
            ocr_text = pytesseract.image_to_string(
                processed_img,
                lang=ocr_lang,
                config="--oem 3 --psm 6"
            )
            if ocr_text.strip():
                text_parts.append(ocr_text)
        except Exception as e:
            raise RuntimeError(f"OCR failed on a PDF page. Error: {e}")

    return "\n".join(text_parts).strip()


def get_text_from_pdf_with_fallback(file_path: str, source_lang: str = "english") -> str:
    text = extract_text_from_pdf(file_path)
    if text.strip():
        return text

    return extract_text_from_pdf_ocr(file_path, source_lang)


def translate_long_text(text: str, target_lang: str) -> str:
    if not text.strip():
        raise ValueError("No readable text found in file.")

    chunk_size = 3000
    chunks = [text[i:i + chunk_size] for i in range(0, len(text), chunk_size)]

    translated_chunks = []
    translator = GoogleTranslator(source="auto", target=target_lang)

    for chunk in chunks:
        translated = translator.translate(chunk)
        translated_chunks.append(translated if translated else "")

    return "\n".join(translated_chunks)


def write_docx(text: str, output_path: str) -> None:
    doc = Document()

    for line in text.splitlines():
        para = doc.add_paragraph()
        run = para.add_run(line)
        run.font.size = Pt(12)

    doc.save(output_path)


def register_font() -> str:
    font_candidates = [
        BASE_DIR / "fonts" / "DejaVuSans.ttf",
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
        Path("/usr/share/fonts/dejavu/DejaVuSans.ttf"),
        Path("C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/arialuni.ttf"),
        Path("C:/Windows/Fonts/Nirmala.ttf"),
        Path("C:/Windows/Fonts/mangal.ttf"),
        Path("C:/Windows/Fonts/shruti.ttf"),
    ]

    for font_path in font_candidates:
        if font_path.exists():
            try:
                pdfmetrics.registerFont(TTFont("CustomUnicode", str(font_path)))
                return "CustomUnicode"
            except Exception:
                continue

    return "Helvetica"


def write_pdf(text: str, output_path: str) -> None:
    font_name = register_font()

    pdf = canvas.Canvas(output_path, pagesize=A4)
    _, height = A4

    left_margin = 40
    top_margin = height - 50
    bottom_margin = 50
    line_height = 16
    max_chars = 90

    pdf.setTitle("Translated PDF")
    pdf.setFont(font_name, 11)

    y = top_margin

    for paragraph in text.splitlines():
        wrapped_lines = textwrap.wrap(paragraph, width=max_chars) or [""]

        for line in wrapped_lines:
            if y < bottom_margin:
                pdf.showPage()
                pdf.setFont(font_name, 11)
                y = top_margin

            pdf.drawString(left_margin, y, line)
            y -= line_height

        y -= 4

    pdf.save()


@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        if "file" not in request.files:
            flash("No file selected.")
            return redirect(url_for("index"))

        file = request.files["file"]
        target_lang = request.form.get("language", "en").strip().lower()
        output_format = request.form.get("format", "pdf").strip().lower()
        source_ocr_language = request.form.get("source_ocr_language", "english").strip().lower()

        if file.filename == "":
            flash("Please choose a file.")
            return redirect(url_for("index"))

        if not allowed_file(file.filename):
            flash("Only PDF and DOCX files are allowed.")
            return redirect(url_for("index"))

        if output_format not in {"pdf", "docx"}:
            flash("Invalid output format.")
            return redirect(url_for("index"))

        filename = secure_filename(file.filename)
        input_path = UPLOAD_FOLDER / filename
        file.save(input_path)

        try:
            ext = input_path.suffix.lower()

            if ext == ".pdf":
                extracted_text = get_text_from_pdf_with_fallback(str(input_path), source_ocr_language)
            else:
                extracted_text = extract_text_from_docx(str(input_path))

            if not extracted_text.strip():
                raise ValueError("PDF me readable text nahi mila. OCR ke baad bhi text extract nahi hua.")

            translated_text = translate_long_text(extracted_text, target_lang)

            output_name = f"translated_{input_path.stem}.{output_format}"
            output_path = OUTPUT_FOLDER / output_name

            if output_format == "pdf":
                write_pdf(translated_text, str(output_path))
            else:
                write_docx(translated_text, str(output_path))

            return render_template(
                "result.html",
                filename=output_name,
                preview=translated_text[:3000]
            )

        except Exception as e:
            flash(f"Error: {str(e)}")
            return redirect(url_for("index"))

    return render_template("index.html")


@app.route("/download/<filename>")
def download_file(filename):
    file_path = OUTPUT_FOLDER / filename

    if not file_path.exists():
        flash("File not found.")
        return redirect(url_for("index"))

    return send_file(str(file_path), as_attachment=True)


if __name__ == "__main__":
    print("Starting Flask app...")
    app.run(host="127.0.0.1", port=5000, debug=False, use_reloader=False)